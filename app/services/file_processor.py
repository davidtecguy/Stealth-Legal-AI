import os
import hashlib
from typing import Optional, Dict, Any
import PyPDF2
from docx import Document as DocxDocument
import tempfile
import shutil
from pathlib import Path

class FileProcessor:
    """Service for processing different file types and extracting text content."""
    
    SUPPORTED_TYPES = {
        'pdf': 'application/pdf',
        'doc': 'application/msword',
        'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    }
    
    def __init__(self, upload_dir: str = "uploads"):
        self.upload_dir = Path(upload_dir)
        self.upload_dir.mkdir(exist_ok=True)
    
    def process_upload(self, file, title: str) -> Dict[str, Any]:
        """Process uploaded file and return document metadata."""
        # Validate file type
        file_extension = self._get_file_extension(file.filename)
        if file_extension not in self.SUPPORTED_TYPES:
            raise ValueError(f"Unsupported file type: {file_extension}")
        
        # Generate unique filename
        unique_filename = self._generate_unique_filename(file.filename)
        file_path = self.upload_dir / unique_filename
        
        # Save file
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        
        # Get file size
        file_size = os.path.getsize(file_path)
        
        # Extract text content
        content = self.extract_text(file_path, file_extension)
        
        # Generate content hash
        content_hash = hashlib.md5(content.encode()).hexdigest() if content else None
        
        return {
            'filename': file.filename,
            'file_path': str(file_path),
            'file_type': file_extension,
            'file_size': file_size,
            'content': content,
            'content_hash': content_hash
        }
    
    def extract_text(self, file_path: str, file_type: str) -> Optional[str]:
        """Extract text content from different file types."""
        try:
            if file_type == 'pdf':
                return self._extract_pdf_text(file_path)
            elif file_type == 'doc':
                return self._extract_doc_text(file_path)
            elif file_type == 'docx':
                return self._extract_docx_text(file_path)
            else:
                return None
        except Exception as e:
            print(f"Error extracting text from {file_path}: {e}")
            return None
    
    def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF file."""
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        return text.strip()
    
    def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            text = []
            for paragraph in doc.paragraphs:
                text.append(paragraph.text)
            return '\n'.join(text)
        except ImportError:
            raise ValueError("python-docx package not installed. Install with: pip install python-docx")
        except Exception as e:
            raise ValueError(f"Failed to extract text from DOCX file: {e}")

    def _extract_doc_text(self, file_path: str) -> str:
        """Extract text from DOC file using antiword or similar tool."""
        # Note: This requires antiword to be installed on the system
        # For now, return a placeholder - in production you'd use antiword or python-docx2txt
        try:
            import subprocess
            result = subprocess.run(['antiword', file_path], capture_output=True, text=True)
            if result.returncode == 0:
                return result.stdout.strip()
            else:
                return f"[DOC file - text extraction not available: {result.stderr}]"
        except (ImportError, FileNotFoundError):
            return "[DOC file - antiword not available for text extraction]"
    
    def _get_file_extension(self, filename: str) -> str:
        """Get file extension from filename."""
        return filename.lower().split('.')[-1] if '.' in filename else ''
    
    def _generate_unique_filename(self, original_filename: str) -> str:
        """Generate unique filename to avoid conflicts."""
        import uuid
        name, ext = os.path.splitext(original_filename)
        unique_id = str(uuid.uuid4())[:8]
        return f"{name}_{unique_id}{ext}"
    
    def get_file_metadata(self, file_path: str) -> Dict[str, Any]:
        """Get additional metadata from file."""
        metadata = {}
        
        try:
            if file_path.lower().endswith('.pdf'):
                metadata.update(self._get_pdf_metadata(file_path))
            elif file_path.lower().endswith('.docx'):
                metadata.update(self._get_docx_metadata(file_path))
        except Exception as e:
            print(f"Error getting metadata from {file_path}: {e}")
        
        return metadata
    
    def _get_pdf_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract metadata from PDF file."""
        metadata = {}
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                if pdf_reader.metadata:
                    metadata['page_count'] = len(pdf_reader.pages)
                    metadata['created_date'] = pdf_reader.metadata.get('/CreationDate')
                    metadata['modified_date'] = pdf_reader.metadata.get('/ModDate')
        except Exception as e:
            print(f"Error extracting PDF metadata: {e}")
        
        return metadata
    
    def _get_docx_metadata(self, file_path: str) -> Dict[str, Any]:
        """Get metadata from DOCX file"""
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            return {
                'title': doc.core_properties.title or '',
                'author': doc.core_properties.author or '',
                'subject': doc.core_properties.subject or '',
                'keywords': doc.core_properties.keywords or '',
                'created': doc.core_properties.created,
                'modified': doc.core_properties.modified,
                'revision': doc.core_properties.revision or 0
            }
        except ImportError:
            return {}
        except Exception:
            return {}

    def _get_doc_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract metadata from DOC file."""
        # Note: This requires antiword to be installed on the system
        # For now, return a placeholder - in production you'd use antiword or python-docx2txt
        try:
            import subprocess
            result = subprocess.run(['antiword', file_path], capture_output=True, text=True)
            if result.returncode == 0:
                return {'title': '', 'author': '', 'subject': '', 'keywords': '', 'created': None, 'modified': None, 'revision': 0}
            else:
                return {'title': '', 'author': '', 'subject': '', 'keywords': '', 'created': None, 'modified': None, 'revision': 0}
        except (ImportError, FileNotFoundError):
            return {'title': '', 'author': '', 'subject': '', 'keywords': '', 'created': None, 'modified': None, 'revision': 0}
    
    def cleanup_file(self, file_path: str):
        """Remove uploaded file."""
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
        except Exception as e:
            print(f"Error cleaning up file {file_path}: {e}")
    
    def get_supported_types(self) -> Dict[str, str]:
        """Get list of supported file types and their MIME types."""
        return self.SUPPORTED_TYPES.copy()
